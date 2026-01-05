import os
import re
import json
import html
import hashlib
from datetime import datetime
from typing import Any, Dict, List
from io import BytesIO

import streamlit as st
import streamlit.components.v1 as components
from dotenv import load_dotenv
from openai import OpenAI

from utils.cv_extract import extract_cv_text


from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from docx import Document





def now_ts() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def stable_id(*parts: str) -> str:
    raw = "||".join([p or "" for p in parts]).encode("utf-8", errors="ignore")
    return hashlib.sha1(raw).hexdigest()[:12]


def safe_int(x: Any, default: int = 0) -> int:
    try:
        return int(x)
    except Exception:
        return default


def init_state():
    if "history" not in st.session_state:
        st.session_state.history = []
    if "ranking_results" not in st.session_state:
        st.session_state.ranking_results = []
    if "selected_id" not in st.session_state:
        st.session_state.selected_id = None
    if "job_text" not in st.session_state:
        st.session_state.job_text = ""
    if "cv_text_paste" not in st.session_state:
        st.session_state.cv_text_paste = ""
    if "shortlist_threshold" not in st.session_state:
        st.session_state.shortlist_threshold = 75
    if "compare_ids" not in st.session_state:
        st.session_state.compare_ids = []

    
    if "uploader_key" not in st.session_state:
        st.session_state.uploader_key = "cv_uploader_0"


def reset_all(clear_history: bool = False, reset_uploads: bool = False):
    st.session_state.selected_id = None
    st.session_state.ranking_results = []
    st.session_state.job_text = ""
    st.session_state.cv_text_paste = ""
    st.session_state.compare_ids = []

    if clear_history:
        st.session_state.history = []

    
    if reset_uploads:
        old = st.session_state.get("uploader_key", "cv_uploader_0")
        try:
            n = int(old.split("_")[-1]) + 1
        except Exception:
            n = 1
        st.session_state.uploader_key = f"cv_uploader_{n}"


def copy_to_clipboard_button(text: str, button_label: str = "Copy to clipboard"):
    escaped = html.escape(text or "")
    components.html(
        f"""
        <button onclick="copyText()" style="
            width:100%;
            padding:0.65rem 1rem;
            border-radius:12px;
            border:none;
            font-weight:800;
            cursor:pointer;
        ">
            {html.escape(button_label)}
        </button>

        <textarea id="copyArea" style="position:absolute; left:-1000px; top:-1000px;">{escaped}</textarea>

        <script>
        function copyText() {{
            var copyText = document.getElementById("copyArea");
            copyText.select();
            copyText.setSelectionRange(0, 999999);
            navigator.clipboard.writeText(copyText.value);
        }}
        </script>
        """,
        height=58,
    )


def render_bullets(items: List[str]):
    if not items:
        st.write("—")
        return
    for x in items:
        if x:
            st.markdown(f"- {x}")


def badge_row(items: List[str], limit: int = 30):
    if not items:
        st.write("—")
        return
    badges = "".join([f"<span class='badge'>{html.escape(s)}</span>" for s in items[:limit] if s])
    st.markdown(badges, unsafe_allow_html=True)





def build_full_text_with_notes(entry: Dict[str, Any]) -> str:
    base = (entry.get("report_text") or "").strip()
    notes = (entry.get("recruiter_notes") or "").strip()

    if notes:
        return f"{base}\n\n## Recruiter notes\n- {notes}\n"
    return base





def make_pdf_bytes(title: str, subtitle: str, body_text: str) -> bytes:
    buff = BytesIO()
    c = canvas.Canvas(buff, pagesize=A4)
    width, height = A4

    left = 2.0 * cm
    top = height - 2.0 * cm
    y = top

    def draw_wrapped(text: str, font="Helvetica", size=10):
        nonlocal y
        c.setFont(font, size)
        max_w = width - 2 * left

        words = text.split()
        line = ""
        for w in words:
            test = (line + " " + w).strip()
            if c.stringWidth(test, font, size) <= max_w:
                line = test
            else:
                c.drawString(left, y, line)
                y -= 12
                line = w
                if y < 2 * cm:
                    c.showPage()
                    y = top
                    c.setFont(font, size)

        if line:
            c.drawString(left, y, line)
            y -= 12
            if y < 2 * cm:
                c.showPage()
                y = top

    c.setFont("Helvetica-Bold", 16)
    c.drawString(left, y, title)
    y -= 18

    if subtitle:
        c.setFont("Helvetica", 10)
        c.drawString(left, y, subtitle[:120])
        y -= 18

    y -= 6
    c.line(left, y, width - left, y)
    y -= 18

    for raw in body_text.replace("\r\n", "\n").splitlines():
        line = raw.strip()
        if not line:
            y -= 8
            if y < 2 * cm:
                c.showPage()
                y = top
            continue

        if line.startswith("## "):
            y -= 6
            draw_wrapped(line.replace("## ", ""), font="Helvetica-Bold", size=12)
            y -= 2
        else:
            if line.startswith(("-", "*")):
                line = "• " + line[1:].strip()
            draw_wrapped(line, font="Helvetica", size=10)

    c.save()
    buff.seek(0)
    return buff.read()


def make_docx_bytes(title: str, subtitle: str, body_text: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    if subtitle:
        doc.add_paragraph(subtitle)

    for raw in body_text.replace("\r\n", "\n").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith(("-", "*")):
            doc.add_paragraph(line[1:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()





def json_to_markdown_report(cv_source: str, data: Dict[str, Any]) -> str:
    score = data.get("overall_score", "N/A")
    reco = data.get("recommendation", "N/A")
    subs = data.get("subscores", {}) or {}
    expl = data.get("explainable_score", {}) or {}
    interview = data.get("interview_guide", {}) or {}

    def bullets(arr):
        arr = arr or []
        return "\n".join([f"- {x}" for x in arr if x]) or "—"

    report = f"""## Candidate
- Source: {cv_source}

## Overall match score
{score}

## Recommendation
{reco}

## Explainable score
**Subscores**
- Skills: {subs.get("skills", 0)}
- Experience: {subs.get("experience", 0)}
- Tools: {subs.get("tools", 0)}
- Domain: {subs.get("domain", 0)}

**Why this score**
{bullets(expl.get("why_this_score", []))}

**Top evidence**
{bullets(expl.get("top_evidence", []))}

## Executive summary
{data.get("summary","") or "—"}

## Key strengths
{bullets(data.get("strengths", []))}

## Gaps & risks
{bullets(data.get("gaps_risks", []))}

## Missing keywords / requirements
{bullets(data.get("missing_keywords", []))}

## Interview guide (focused on gaps)
**Critical**
{bullets(interview.get("critical", []))}

**Nice-to-have**
{bullets(interview.get("nice_to_have", []))}

## CV improvement suggestions
{bullets(data.get("cv_improvements", []))}
"""
    return report.strip()





def build_schema_instruction() -> str:
    return """
Return ONLY a valid JSON object with EXACTLY these keys:
{
  "overall_score": <integer 0-100>,
  "recommendation": <one of "Strong Yes","Yes","Maybe","No">,
  "subscores": {
    "skills": <0-100>,
    "experience": <0-100>,
    "tools": <0-100>,
    "domain": <0-100>
  },
  "explainable_score": {
    "why_this_score": [<bullet strings>],
    "top_evidence": [<short evidence strings referencing CV facts>]
  },
  "strengths": [<bullet strings>],
  "gaps_risks": [<bullet strings>],
  "missing_keywords": [<strings>],
  "interview_guide": {
    "critical": [<questions/checks>],
    "nice_to_have": [<questions/checks>]
  },
  "cv_improvements": [<bullet strings>],
  "summary": <1-2 sentence executive summary>
}

Rules:
- Do NOT invent. Only claim what is explicitly in the CV.
- Keep bullets short and business-ready.
""".strip()


def call_openai_json(client: OpenAI, model: str, job_text: str, cv_text: str) -> Dict[str, Any]:
    prompt = f"""
You are a senior HR Talent Intelligence analyst.

Task: Assess CV vs Job Description with a factual, explainable evaluation.
{build_schema_instruction()}

JOB DESCRIPTION:
{job_text}

CANDIDATE CV:
{cv_text}
""".strip()

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "Be rigorous, factual, and concise. Output JSON only."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            response_format={"type": "json_object"},
        )
        return json.loads(resp.choices[0].message.content)
    except Exception:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "Be rigorous, factual, and concise. Output JSON only."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
        )
        content = resp.choices[0].message.content or ""
        m = re.search(r"\{.*\}", content, flags=re.DOTALL)
        if not m:
            raise ValueError("Model did not return JSON.")
        return json.loads(m.group(0))


def update_notes(entry_id: str, notes: str):
    for h in st.session_state.history:
        if h["id"] == entry_id:
            h["recruiter_notes"] = notes
            break
    for r in st.session_state.ranking_results:
        if r["id"] == entry_id:
            r["recruiter_notes"] = notes
            break





load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    st.error("OPENAI_API_KEY not found in .env")
    st.stop()

MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
client = OpenAI(api_key=api_key)

st.set_page_config(page_title="Talent Match Assistant", page_icon="🧠", layout="wide")
init_state()




st.markdown(
    """
    <style>
      .block-container { padding-top: 3.2rem; }

      .tma-card {
        border: 1px solid rgba(255,255,255,0.08);
        border-radius: 22px;
        padding: 22px 26px;
        background: rgba(255,255,255,0.08);
        backdrop-filter: blur(10px);
        min-height: 104px;
        box-shadow: 0 14px 36px rgba(0,0,0,0.35);
      }
      .tma-title {
        font-size: 1.95rem;
        font-weight: 800;
        margin: 0;
        line-height: 1.25;
        color: rgba(255,255,255,0.92);
      }

      .stTextArea textarea, .stFileUploader, .stButton button, .stSelectbox, .stMultiSelect {
        border-radius: 14px !important;
      }

      .kpi-big { font-size: 2.6rem; font-weight: 900; color: rgba(255,255,255,0.92); }
      .kpi-mid { font-size: 2.1rem; font-weight: 850; color: rgba(255,255,255,0.92); }
      .kpi-label { opacity: 0.72; color: rgba(255,255,255,0.72); }

      .badge {
        display:inline-block;
        padding: 6px 10px;
        border-radius: 999px;
        border: 1px solid rgba(255,255,255,0.14);
        background: rgba(255,255,255,0.06);
        margin: 4px 6px 0 0;
        font-size: 0.85rem;
        opacity: 0.95;
      }

      .mini { opacity: 0.75; font-size: 0.92rem; }
    </style>
    """,
    unsafe_allow_html=True,
)


st.markdown(
    """
    <div class="tma-card">
      <p class="tma-title">Talent Match Assistant</p>
      <div class="mini">Shortlist • Ranking • Explainable scoring • Interview guide • Notes</div>
    </div>
    """,
    unsafe_allow_html=True,
)
st.write("")





with st.sidebar:
    st.markdown("### Controls")

    c1, c2 = st.columns(2)
    with c1:
        if st.button("Clear Inputs", use_container_width=True):
            
            reset_all(clear_history=False, reset_uploads=False)
            st.rerun()
    with c2:
        if st.button("Reset Session", use_container_width=True):
            
            reset_all(clear_history=True, reset_uploads=True)
            st.rerun()

    st.session_state.shortlist_threshold = st.slider(
        "Shortlist threshold",
        min_value=0,
        max_value=100,
        value=st.session_state.shortlist_threshold,
        step=1,
    )

    
    if st.session_state.selected_id:
        selected = next((h for h in st.session_state.history if h["id"] == st.session_state.selected_id), None)
        if selected:
            st.markdown("#### Actions")

            full_text = build_full_text_with_notes(selected)

            
            copy_to_clipboard_button(full_text, "Copy selected report")

            subtitle = f"Score: {selected.get('overall_score')}/100 | Recommendation: {selected.get('recommendation')} | {selected.get('cv_source')}"

            
            pdf_bytes = make_pdf_bytes(
                title="Talent Match Assistant",
                subtitle=subtitle,
                body_text=full_text,
            )
            st.download_button(
                "Download PDF",
                data=pdf_bytes,
                file_name=f"talent_match_{(selected.get('cv_source','candidate')).replace(' ','_')}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )

            
            docx_bytes = make_docx_bytes(
                title="Talent Match Assistant",
                subtitle=subtitle,
                body_text=full_text,
            )
            st.download_button(
                "Download DOCX",
                data=docx_bytes,
                file_name=f"talent_match_{(selected.get('cv_source','candidate')).replace(' ','_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

    st.markdown("---")
    st.markdown("### History")
    if st.session_state.history:
        labels = [
            f"{i+1}. {h['timestamp']} • {h['overall_score']}/100 • {h['recommendation']} • {h['cv_source']}"
            for i, h in enumerate(st.session_state.history)
        ]
        sel_label = st.selectbox("Select analysis", labels, index=len(labels) - 1)
        idx = labels.index(sel_label)
        if st.button("Load selected", use_container_width=True):
            st.session_state.selected_id = st.session_state.history[idx]["id"]
            st.rerun()
    else:
        st.caption("No analyses yet.")





col1, col2 = st.columns(2, gap="large")

with col1:
    st.subheader("Job Description")
    st.session_state.job_text = st.text_area(
        "Paste the job description",
        value=st.session_state.job_text,
        height=320,
        placeholder="Responsibilities, requirements, stack..."
    )

with col2:
    st.subheader("Candidate CV(s)")
    cv_files = st.file_uploader(
        "Upload CV(s) (PDF or DOCX) — multiple selection supported",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        key=st.session_state.uploader_key,  
    )
    st.caption("Fallback: paste CV text below (used only if no files uploaded).")
    st.session_state.cv_text_paste = st.text_area(
        "Paste CV text (optional)",
        value=st.session_state.cv_text_paste,
        height=160,
    )

st.write("")
analyze = st.button("Analyze match", type="primary")


def add_report_to_history(entry: Dict[str, Any]):
    st.session_state.history.append(entry)
    st.session_state.selected_id = entry["id"]





if analyze:
    job_text = st.session_state.job_text.strip()
    if not job_text:
        st.error("Please paste a Job Description.")
        st.stop()

    results: List[Dict[str, Any]] = []

    if cv_files:
        with st.spinner(f"Analyzing {len(cv_files)} CV(s)..."):
            for f in cv_files:
                try:
                    file_bytes = f.read()
                    cv_text, detected = extract_cv_text(f.name, file_bytes)
                    cv_source = f"{f.name} ({detected.upper()})"
                except Exception as e:
                    rid = stable_id(job_text[:120], f.name, "extract_error")
                    results.append({
                        "id": rid,
                        "timestamp": now_ts(),
                        "cv_source": f.name,
                        "overall_score": 0,
                        "recommendation": "No",
                        "subscores": {"skills": 0, "experience": 0, "tools": 0, "domain": 0},
                        "summary": "Extraction failed.",
                        "missing_keywords": [],
                        "strengths": [],
                        "gaps_risks": [f"Failed to extract CV text: {e}"],
                        "interview_guide": {"critical": ["Re-upload as DOCX or paste text"], "nice_to_have": []},
                        "cv_improvements": [],
                        "explainable_score": {"why_this_score": ["No text extracted"], "top_evidence": []},
                        "recruiter_notes": "",
                        "report_text": "Extraction failed.",
                    })
                    continue

                data = call_openai_json(client, MODEL, job_text, cv_text)

                score = safe_int(data.get("overall_score"), 0)
                reco = (data.get("recommendation") or "Maybe").strip()
                subs = data.get("subscores") or {}
                interview = data.get("interview_guide") or {"critical": [], "nice_to_have": []}

                rid = stable_id(job_text[:160], cv_source, str(score), reco)
                report_text = json_to_markdown_report(cv_source, data)

                results.append({
                    "id": rid,
                    "timestamp": now_ts(),
                    "cv_source": cv_source,
                    "overall_score": score,
                    "recommendation": reco,
                    "subscores": {
                        "skills": safe_int(subs.get("skills"), 0),
                        "experience": safe_int(subs.get("experience"), 0),
                        "tools": safe_int(subs.get("tools"), 0),
                        "domain": safe_int(subs.get("domain"), 0),
                    },
                    "summary": data.get("summary", ""),
                    "missing_keywords": data.get("missing_keywords") or [],
                    "strengths": data.get("strengths") or [],
                    "gaps_risks": data.get("gaps_risks") or [],
                    "interview_guide": {
                        "critical": (interview.get("critical") or []),
                        "nice_to_have": (interview.get("nice_to_have") or []),
                    },
                    "cv_improvements": data.get("cv_improvements") or [],
                    "explainable_score": data.get("explainable_score") or {"why_this_score": [], "top_evidence": []},
                    "recruiter_notes": "",
                    "report_text": report_text,
                })

        results.sort(key=lambda r: safe_int(r["overall_score"]), reverse=True)
        st.session_state.ranking_results = results

        for r in results:
            add_report_to_history(r)

        st.rerun()

    else:
        cv_text = st.session_state.cv_text_paste.strip()
        if not cv_text:
            st.error("Upload at least one CV OR paste CV text.")
            st.stop()

        with st.spinner("Analyzing pasted CV text..."):
            data = call_openai_json(client, MODEL, job_text, cv_text)

        score = safe_int(data.get("overall_score"), 0)
        reco = (data.get("recommendation") or "Maybe").strip()
        subs = data.get("subscores") or {}
        interview = data.get("interview_guide") or {"critical": [], "nice_to_have": []}

        rid = stable_id(job_text[:160], "pasted_text", str(score), reco)
        report_text = json_to_markdown_report("Pasted text", data)

        entry = {
            "id": rid,
            "timestamp": now_ts(),
            "cv_source": "Pasted text",
            "overall_score": score,
            "recommendation": reco,
            "subscores": {
                "skills": safe_int(subs.get("skills"), 0),
                "experience": safe_int(subs.get("experience"), 0),
                "tools": safe_int(subs.get("tools"), 0),
                "domain": safe_int(subs.get("domain"), 0),
            },
            "summary": data.get("summary", ""),
            "missing_keywords": data.get("missing_keywords") or [],
            "strengths": data.get("strengths") or [],
            "gaps_risks": data.get("gaps_risks") or [],
            "interview_guide": {
                "critical": (interview.get("critical") or []),
                "nice_to_have": (interview.get("nice_to_have") or []),
            },
            "cv_improvements": data.get("cv_improvements") or [],
            "explainable_score": data.get("explainable_score") or {"why_this_score": [], "top_evidence": []},
            "recruiter_notes": "",
            "report_text": report_text,
        }

        st.session_state.ranking_results = [entry]
        add_report_to_history(entry)
        st.rerun()





if st.session_state.ranking_results:
    st.divider()

    st.subheader("Shortlist")
    thr = st.session_state.shortlist_threshold
    shortlist = [r for r in st.session_state.ranking_results if safe_int(r["overall_score"]) >= thr]
    st.caption(f"Auto-shortlist: candidates with score ≥ {thr}")

    if shortlist:
        st.dataframe(
            [
                {
                    "Score": r["overall_score"],
                    "Recommendation": r["recommendation"],
                    "Candidate": r["cv_source"],
                    "Missing keywords (count)": len(r.get("missing_keywords", [])),
                }
                for r in shortlist
            ],
            use_container_width=True,
            hide_index=True,
        )
    else:
        st.info("No candidates meet the shortlist threshold yet.")

    st.subheader("Ranking & comparison")
    st.caption("Select 2–5 candidates to compare side-by-side.")

    options = [(r["id"], f"{r['overall_score']}/100 • {r['recommendation']} • {r['cv_source']}") for r in st.session_state.ranking_results]
    id_to_label = {i: lbl for i, lbl in options}

    st.session_state.compare_ids = st.multiselect(
        "Compare candidates",
        options=[i for i, _ in options],
        format_func=lambda x: id_to_label.get(x, x),
        default=st.session_state.compare_ids[:5],
        max_selections=5,
    )

    open_id = st.selectbox(
        "Open candidate details",
        options=[i for i, _ in options],
        format_func=lambda x: id_to_label.get(x, x),
        index=0,
    )
    if st.button("Open selected candidate"):
        st.session_state.selected_id = open_id
        st.rerun()

    if len(st.session_state.compare_ids) >= 2:
        compare = [next(r for r in st.session_state.ranking_results if r["id"] == cid) for cid in st.session_state.compare_ids]
        cols = st.columns(len(compare), gap="large")
        for c, r in zip(cols, compare):
            with c:
                st.markdown(f"**{r['cv_source']}**")
                st.markdown(f"<div class='kpi-big'>{r['overall_score']}/100</div>", unsafe_allow_html=True)
                st.caption(r["recommendation"])
                st.markdown("**Top strengths**")
                render_bullets((r.get("strengths") or [])[:5])
                st.markdown("**Top gaps**")
                render_bullets((r.get("gaps_risks") or [])[:5])
                st.markdown("**Missing keywords**")
                badge_row((r.get("missing_keywords") or [])[:12], limit=12)





if st.session_state.selected_id:
    sel = next((h for h in st.session_state.history if h["id"] == st.session_state.selected_id), None)
    if sel:
        st.divider()
        st.subheader("Candidate details")

        k1, k2, k3 = st.columns(3, gap="large")
        with k1:
            st.markdown(
                f"<div class='tma-card' style='text-align:center; min-height:140px;'><div class='kpi-big'>{sel['overall_score']}/100</div><div class='kpi-label'>Overall</div></div>",
                unsafe_allow_html=True,
            )
        with k2:
            st.markdown(
                f"<div class='tma-card' style='text-align:center; min-height:140px;'><div class='kpi-mid'>{html.escape(sel['recommendation'])}</div><div class='kpi-label'>Recommendation</div></div>",
                unsafe_allow_html=True,
            )
        with k3:
            subs = sel.get("subscores", {})
            st.markdown(
                f"""
                <div class='tma-card' style='min-height:140px;'>
                  <div style='font-weight:900;margin-bottom:6px;'>Explainable score</div>
                  <div class='mini'>Skills {safe_int(subs.get('skills'))} • Exp {safe_int(subs.get('experience'))}</div>
                  <div class='mini'>Tools {safe_int(subs.get('tools'))} • Domain {safe_int(subs.get('domain'))}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

        st.write("")
        st.markdown("#### Missing keywords / requirements")
        badge_row(sel.get("missing_keywords") or [], limit=30)

        st.markdown("#### Interview guide (focused on gaps)")
        ig = sel.get("interview_guide") or {"critical": [], "nice_to_have": []}
        a, b = st.columns(2, gap="large")
        with a:
            st.markdown("**Critical checks**")
            render_bullets(ig.get("critical") or [])
        with b:
            st.markdown("**Nice-to-have**")
            render_bullets(ig.get("nice_to_have") or [])

        st.markdown("#### Recruiter notes")
        notes_key = f"notes_{sel['id']}"
        if notes_key not in st.session_state:
            st.session_state[notes_key] = sel.get("recruiter_notes", "")

        new_notes = st.text_area(
            "Write notes (saved automatically in this session)",
            value=st.session_state[notes_key],
            height=140,
        )
        if new_notes != st.session_state[notes_key]:
            st.session_state[notes_key] = new_notes
            update_notes(sel["id"], new_notes)

        st.markdown("#### Full report")
        st.markdown(sel.get("report_text", "—"))

